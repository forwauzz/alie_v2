"""Word export (PRD §4.1, §8.1).

Markdown is for machines. The chronology is a document a paralegal reads, edits and hands
to opposing counsel, and that means Word.

Two documents, because they answer different questions:

- **The chronology** — what the engine produced, in the three-column shape the firm uses,
  with every locator intact. Rows held back by a rule appear in their own section rather
  than vanishing (§10.1).
- **The gap** — the same chronology set beside the answer key: what the key has and the
  engine missed, what the engine produced that the key does not, and where the dates
  disagree. This is the document that says whether the thing works.

Nothing here decides anything. Both documents render what the stores already hold; the
gap report scores against a gold and shows its working rather than a single number.
"""

from __future__ import annotations

import re
import sqlite3
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..models import Row
from ..packs import Pack
from ..packs import load as load_pack
from ..stores import cases
from . import render as render_stage

#: Muted grey for the locator column and provenance notes — present, never competing with
#: the clinical content.
MUTED = RGBColor(0x66, 0x66, 0x66)
FLAG = RGBColor(0xB0, 0x3A, 0x2B)


#: What XML 1.0 *permits*, as a whitelist. Blacklisting the obvious control characters was
#: not enough — real scans also carry unpaired surrogates and C1 bytes that lxml rejects,
#: and each one only surfaced as another crash three minutes into a 300-page run. Naming
#: the allowed set is the only version of this that terminates.
_XML_ALLOWED = re.compile("[^\x09\x0a\x0d\x20-\ud7ff\ue000-\ufffd\U00010000-\U0010ffff]")


def xml_safe(text: str) -> str:
    """Drop characters Word cannot represent.

    Real scans are full of them: a failed OCR pass emits `\\rLllll\\{vÊ`, and page-break
    bytes ride along with fax headers.

    This changes only what is *displayed*. The record keeps its original text and its span,
    so the citation still resolves to the exact source bytes — a paralegal clicking through
    lands on the same characters whatever the document could draw. Dropping an unprintable
    byte from a rendering is not the same as altering the evidence.
    """
    return _XML_ALLOWED.sub("", text)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(xml_safe(text), level=level)


def _note(doc: Document, text: str, colour: RGBColor = MUTED) -> None:
    para = doc.add_paragraph()
    run = para.add_run(xml_safe(text))
    run.font.size = Pt(9)
    run.font.color.rgb = colour


def _chronology_table(doc: Document, rows: list[Row], pack: Pack, folders: dict[str, str],
                      *, show_code: bool) -> None:
    headers = [c["header"] for c in pack.output.get("columns", [])] or [
        "Date", "Document", "Contenu",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers, strict=False):
        cell.text = ""
        run = cell.paragraphs[0].add_run(xml_safe(header))
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = xml_safe(row.row_date.render()) if not row.is_undated else "—"

        # Locators, one per contributing bundle. Both are retained on a cross-bundle
        # union, because "which page" is the whole point of a citation (§8.1).
        locator = cells[1].paragraphs[0]
        for index, citation in enumerate(row.locators):
            if index:
                locator.add_run("\n")
            text = render_stage.locator_text(
                pack, folders.get(citation.bundle_id, citation.bundle_id),
                citation.display_page, citation.needs_flag,
            )
            run = locator.add_run(xml_safe(text))
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED

        content = cells[2].paragraphs[0]
        code = pack.class_code(row.doc_class) if show_code else None
        title = content.add_run(xml_safe(f"[{code}] {row.title}" if code else row.title))
        title.bold = True
        if row.illegible_reason:
            content.add_run(xml_safe(f"\nIllisible — {row.illegible_reason}")).italic = True
        for bullet in row.bullets:
            content.add_run(xml_safe(f"\n• {bullet.text}"))
        if row.warns:
            warn = content.add_run(f"\n⚠ confiance {row.confidence:.2f}")
            warn.font.color.rgb = FLAG
            warn.font.size = Pt(9)


def to_docx(
    conn: sqlite3.Connection,
    case_id: str,
    rows: list[Row],
    path: Path,
    *,
    flags: dict | None = None,
) -> Path:
    """Write the chronology as a Word document."""
    case = cases.get_case(conn, case_id)
    pack = load_pack(case["primary_pack"])
    folders = {b["id"]: b["folder_label"] for b in cases.bundles_for_case(conn, case_id)}
    show_code = bool((flags or {}).get("render.doctype_code"))

    doc = Document()
    doc.add_heading(xml_safe(f"Chronologie médicale — {case['name']}"), level=0)
    _note(
        doc,
        f"Régime {pack.id.upper()} · pack v{pack.version} · {len(rows)} ligne(s). "
        "Chaque ligne est transcrite du dossier et citée à sa page; rien n'est résumé.",
    )

    kept = [r for r in rows if not r.excluded_by]
    undated = [r for r in kept if r.is_undated]
    dated = [r for r in kept if not r.is_undated]

    if undated:
        # Undated rows lead the document, so they are the first thing reviewed rather than
        # the last thing discovered (§8.5).
        heading = pack.output.get("rows", {}).get(
            "undated_heading", "SANS DATE — {n} documents à dater"
        )
        _heading(doc, heading.format(n=len(undated)), level=1)
        _chronology_table(doc, undated, pack, folders, show_code=show_code)
        _heading(doc, "Chronologie", level=1)

    _chronology_table(doc, dated, pack, folders, show_code=show_code)

    held = [r for r in rows if r.excluded_by]
    if held:
        # Removing pages from a legal record is never destructive on the source (§10.1).
        _heading(doc, "Retiré par règle — conservé au dossier", level=1)
        _note(doc, "Ces documents restent au manifeste. La règle qui les a retirés est nommée.")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, ("Date", "Document", "Règle"),
                               strict=False):
            cell.paragraphs[0].add_run(xml_safe(header)).bold = True
        for row in held:
            cells = table.add_row().cells
            cells[0].text = row.row_date.value.isoformat() if row.row_date.value else "—"
            cells[1].text = xml_safe(row.title)
            cells[2].text = xml_safe(row.excluded_by or "")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def gap_to_docx(
    path: Path,
    *,
    case_name: str,
    rows: list[Row],
    gold: list[dict[str, Any]],
    dates: dict[str, int],
    content: dict[str, Any],
    caveats: list[str],
    header: dict[str, Any] | None = None,
) -> Path:
    """Write the answer-key comparison as a Word document.

    Shows its working. A single recall percentage invites a decision the number cannot
    support — the misses have to be readable, because most of the judgement about whether
    a miss matters is in the wording of the line itself (§11.4).
    """
    doc = Document()
    doc.add_heading(xml_safe(f"Écart vs clé de correction — {case_name}"), level=0)

    for line in caveats:
        _note(doc, line, FLAG)

    if header:
        _heading(doc, "Conditions", level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for key, value in header.items():
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(xml_safe(str(key))).bold = True
            cells[1].text = xml_safe(str(value))

    _heading(doc, "Dates", level=1)
    scored = dates["exact"] + dates["near"] + dates["missing"]
    para = doc.add_paragraph()
    para.add_run(
        f"{dates['exact']}/{scored} exactes · "
        f"{dates['exact'] + dates['near']}/{scored} à 3 jours près · "
        f"{dates['missing']} absentes"
    ).bold = True

    _heading(doc, "Contenu", level=1)
    doc.add_paragraph(
        f"{content['hits']}/{content['total']} lignes de la clé retrouvées "
        f"({content['hits'] / max(content['total'], 1):.0%})."
    )

    if content.get("misses"):
        _heading(doc, "Lignes de la clé non retrouvées", level=2)
        _note(
            doc,
            "À adjuger. Une ligne « manquante » est parfois la reformulation de la "
            "technicienne plutôt qu'un contenu absent du dossier (§11.4).",
        )
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        for cell, head in zip(table.rows[0].cells, ("Date", "Ligne de la clé"),
                             strict=False):
            cell.paragraphs[0].add_run(xml_safe(head)).bold = True
        for miss in content["misses"]:
            when, _, text = miss.partition(" | ")
            cells = table.add_row().cells
            cells[0].text = xml_safe(when)
            cells[1].text = xml_safe(text)

    _heading(doc, "Lignes produites par le moteur", level=1)
    doc.add_paragraph(f"{len(rows)} ligne(s) produite(s) pour {len(gold)} ligne(s) de clé.")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for row in sorted(rows, key=lambda r: (r.row_date.value or date.min)):
        when = row.row_date.value.isoformat() if row.row_date.value else "—"
        para.add_run(f"\n{when}  {row.title}").bold = True
        for bullet in row.bullets[:6]:
            para.add_run(f"\n    • {bullet.text}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
